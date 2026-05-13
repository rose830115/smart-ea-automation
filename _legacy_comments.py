"""Legacy rule-based comment generator (replaced by ChatGPT-based comment_generator.py)."""
from __future__ import annotations
import json
from collections import defaultdict
from pathlib import Path
from statistics import mean
from typing import Any

from openpyxl import Workbook

from smart_ea_automation import (
    ZONE_LABELS,
    average,
    class_label,
    clean_text,
    format_number,
    group_area_values,
    join_list,
    moisture_average,
    numeric_values,
    safe_float,
    style_sheet,
)


def moisture_class_stats(rows: list[dict[str, Any]]) -> dict[str, dict[str, float]]:
    grouped: dict[str, list[float]] = defaultdict(list)
    for row in rows:
        cls = clean_text(row.get("Classification"))
        for idx in range(1, 6):
            value = safe_float(row.get(f"Moisture (%) - {idx}"))
            if value is not None:
                grouped[cls].append(value)
    stats: dict[str, dict[str, float]] = {}
    for cls, values in grouped.items():
        failed = [value for value in values if value > 10]
        stats[cls] = {
            "average": mean(values),
            "maximum": max(values),
            "failed_ratio": len(failed) / len(values) * 100,
            "count": len(values),
        }
    return stats


def max_area(values: dict[str, float]) -> tuple[str, float] | None:
    if not values:
        return None
    return max(values.items(), key=lambda item: item[1])


def zone_metrics(data: dict[str, list[dict[str, Any]]]) -> dict[str, Any]:
    env = data["Env. Data"]
    moisture = data["Moisture"]
    outside = [r for r in env if r.get("Main Zone") == "OUTSIDE"]
    outside_rh = average(numeric_values(outside, "Humidity (%)"))
    outside_pm10 = average(numeric_values(outside, "PM10 (μg/m³)"))

    metrics: dict[str, Any] = {}
    for zone in ["RMW", "PL", "FGW"]:
        zone_env = [r for r in env if r.get("Main Zone") == zone]
        zone_moisture = [r for r in moisture if r.get("Main Zone") == zone]
        rh = average(numeric_values(zone_env, "Humidity (%)"))
        co2 = average(numeric_values(zone_env, "CO2 (ppm)"))
        wind = average(numeric_values(zone_env, "Wind Flow (m/s)"))
        pm10 = average(numeric_values(zone_env, "PM10 (μg/m³)"))
        humidity_permeability = rh / outside_rh * 100 if rh is not None and outside_rh else None
        infiltration = pm10 / outside_pm10 * 100 if pm10 is not None and outside_pm10 else None
        area_rh = group_area_values(zone_env, "Humidity (%)")
        area_co2 = group_area_values(zone_env, "CO2 (ppm)")
        area_wind = group_area_values(zone_env, "Wind Flow (m/s)")
        area_pm10 = group_area_values(zone_env, "PM10 (μg/m³)")
        area_humidity_permeability = {a: v / outside_rh * 100 for a, v in area_rh.items() if outside_rh}
        area_spore_infiltration = {a: v / outside_pm10 * 100 for a, v in area_pm10.items() if outside_pm10}
        failed_objects: dict[str, str] = {}
        for row in zone_moisture:
            avg = moisture_average(row)
            if avg is not None and avg > 10:
                obj = clean_text(row.get("Object"))
                if obj:
                    failed_objects.setdefault(obj.casefold(), obj)
        class_stats = moisture_class_stats(zone_moisture)
        total_moisture_values = sum(int(s["count"]) for s in class_stats.values())
        failed_moisture_values = sum(round(s["failed_ratio"] / 100 * s["count"]) for s in class_stats.values())
        failed_ratio = failed_moisture_values / total_moisture_values * 100 if total_moisture_values else 0
        high_moisture_classes = [
            class_label(cls)
            for cls, stat in sorted(class_stats.items(), key=lambda item: item[1]["failed_ratio"], reverse=True)
            if stat["failed_ratio"] >= 50
        ]
        high_rh_areas = [a for a, v in area_rh.items() if v > 60]
        high_hp_areas = [a for a, v in area_humidity_permeability.items() if v > 90]
        high_co2_areas = [a for a, v in area_co2.items() if v > 600]
        low_wind_areas = [a for a, v in area_wind.items() if v < 0.4]
        high_pm10_areas = [a for a, v in area_spore_infiltration.items() if v > 90]

        score = 0.0
        area_count = max(1, len(area_rh))
        if rh is not None and rh > 60:
            score += min(20, (rh - 60) * 2)
        if humidity_permeability is not None and humidity_permeability > 90:
            score += min(15, (humidity_permeability - 90) * 0.8)
        if high_co2_areas:
            score += len(high_co2_areas) / area_count * 15
        if low_wind_areas:
            score += len(low_wind_areas) / area_count * 15
        if high_pm10_areas:
            score += len(high_pm10_areas) / area_count * 15
        if high_moisture_classes:
            score += min(20, max(stat["failed_ratio"] for stat in class_stats.values()) * 0.2)

        metrics[zone] = {
            "label": ZONE_LABELS[zone],
            "areas": sorted({clean_text(r.get("Report Area")) for r in zone_env if r.get("Report Area")}),
            "area_count": area_count,
            "avg_rh": rh,
            "avg_co2": co2,
            "avg_wind": wind,
            "avg_pm10": pm10,
            "humidity_permeability": humidity_permeability,
            "spore_infiltration": infiltration,
            "moisture_failed_ratio": failed_ratio,
            "moisture_class_stats": class_stats,
            "high_moisture_classes": high_moisture_classes,
            "failed_objects": sorted(failed_objects.values(), key=str.casefold),
            "area_rh": area_rh,
            "area_humidity_permeability": area_humidity_permeability,
            "area_co2": area_co2,
            "area_wind": area_wind,
            "area_spore_infiltration": area_spore_infiltration,
            "high_rh_areas": high_rh_areas,
            "high_hp_areas": high_hp_areas,
            "high_co2_areas": high_co2_areas,
            "low_wind_areas": low_wind_areas,
            "high_pm10_areas": high_pm10_areas,
            "risk_score": round(score, 1),
        }

    metrics["OUTSIDE"] = {
        "avg_rh": outside_rh,
        "avg_pm10": outside_pm10,
        "rows": outside,
    }
    return metrics


def issue_names(m: dict[str, Any]) -> list[str]:
    issues = []
    area_count = max(1, m.get("area_count", len(m.get("areas", [])) or 1))
    if m["avg_rh"] is not None and m["avg_rh"] > 60:
        issues.append("the relative humidity")
    if (m["humidity_permeability"] is not None and m["humidity_permeability"] > 90) or len(m["high_hp_areas"]) / area_count >= 0.5:
        issues.append("the humidity permeability rate")
    if (m["avg_co2"] is not None and m["avg_co2"] > 600) or len(m["high_co2_areas"]) / area_count >= 0.5:
        issues.append("the carbon dioxide level")
    if (m["avg_wind"] is not None and m["avg_wind"] < 0.4) or len(m["low_wind_areas"]) / area_count >= 0.5:
        issues.append("the mold spore blow index")
    if (m["spore_infiltration"] is not None and m["spore_infiltration"] > 90) or len(m["high_pm10_areas"]) / area_count >= 0.5:
        issues.append("the spore infiltration rate")
    if m["moisture_failed_ratio"] > 50 or m["high_moisture_classes"]:
        issues.append("the object moisture content")
    return issues


def zone_comment(zone: str, m: dict[str, Any]) -> str:
    label = m["label"]
    issues = issue_names(m)
    lead = f"The main issue in the {label} was {issues[0]}. " if len(issues) == 1 else f"The main issues in the {label} were {join_list(issues)}. "
    if not issues:
        lead = f"The environmental parameters in the {label} were generally controlled, but continued monitoring is still recommended. "

    sentences = [lead]
    if m["avg_rh"] is not None:
        if m["avg_rh"] > 60:
            max_rh = max_area(m["area_rh"])
            max_text = f", with the highest value in {max_rh[0]} at {format_number(max_rh[1])}%" if max_rh else ""
            areas = f" Areas affected included {join_list(m['high_rh_areas'])}." if m["high_rh_areas"] else ""
            sentences.append(f"The average relative humidity was {format_number(m['avg_rh'])}%{max_text}, exceeding the 60% standard and creating conditions favorable for mold germination.{areas} ")
        elif m["avg_rh"] >= 55:
            sentences.append(f"Although the average relative humidity ({format_number(m['avg_rh'])}%) met the standard, it was close to the 60% threshold, so humidity management should remain in place. ")
    if (m["humidity_permeability"] is not None and m["humidity_permeability"] > 90) or m["high_hp_areas"]:
        extra = " Since the value was higher than 100%, indoor moisture may have been trapped and not effectively released outdoors." if m["humidity_permeability"] > 100 else ""
        areas = f" Affected areas included {join_list(m['high_hp_areas'])}." if m["high_hp_areas"] else ""
        sentences.append(f"The humidity permeability rate was {format_number(m['humidity_permeability'])}%, indicating that outdoor humidity had a strong influence on indoor humidity.{extra}{areas} ")
    if m["moisture_failed_ratio"] > 0 or m["high_moisture_classes"]:
        objects = join_list(m["failed_objects"][:8])
        high_classes = join_list(m["high_moisture_classes"])
        if m["high_moisture_classes"]:
            detail = f" The higher-failure object groups were {high_classes}." if high_classes else ""
            examples = f" Examples included {objects}." if objects else ""
            sentences.append(f"Object moisture content was a mold-risk source because several checked objects exceeded 10%.{detail}{examples} These objects could provide moisture and nutrients for mold growth. ")
        elif m["moisture_failed_ratio"] > 50:
            sentences.append(f"Moisture content in most checked objects exceeded 10% (measured-point failed ratio: {format_number(m['moisture_failed_ratio'])}%), including {objects}; these objects could readily provide moisture and nutrients for mold growth. ")
        else:
            sentences.append(f"Some checked objects exceeded 10% moisture content, including {objects}, and should be treated as localized mold-growth hotspots. ")
    if (m["avg_wind"] is not None and m["avg_wind"] < 0.4) or m["low_wind_areas"]:
        areas = f" Low-wind areas included {join_list(m['low_wind_areas'])}." if m["low_wind_areas"] else ""
        if m["avg_wind"] is not None and m["avg_wind"] < 0.4:
            sentences.append(f"The average wind speed was {format_number(m['avg_wind'], 2)} m/s, below the suggested 0.4 m/s level. Low mold spore blow index may allow spores and dust to settle on object surfaces.{areas} ")
        else:
            sentences.append(f"Although the zone-average wind speed was {format_number(m['avg_wind'], 2)} m/s, low mold spore blow index was still found locally.{areas} Such areas may allow spores and dust to settle on object surfaces. ")
    if (m["spore_infiltration"] is not None and m["spore_infiltration"] > 90) or m["high_pm10_areas"]:
        max_inf = max_area(m["area_spore_infiltration"])
        max_text = f" The highest area-level value was {format_number(max_inf[1])}% in {max_inf[0]}." if max_inf else ""
        areas = f" Affected areas included {join_list(m['high_pm10_areas'])}." if m["high_pm10_areas"] else ""
        sentences.append(f"The estimated spore infiltration rate was {format_number(m['spore_infiltration'])}%, suggesting that outdoor particles could enter or remain indoors and gradually increase airborne spore accumulation.{max_text}{areas} ")
    if (m["avg_co2"] is not None and m["avg_co2"] > 600) or m["high_co2_areas"]:
        areas = f" High-CO2 areas included {join_list(m['high_co2_areas'])}." if m["high_co2_areas"] else ""
        if m["avg_co2"] is not None and m["avg_co2"] > 600:
            sentences.append(f"The average carbon dioxide level was {format_number(m['avg_co2'], 0)} ppm, which was higher than the recommended 600 ppm level and may indicate insufficient air exchange or microbiological activity.{areas} ")
        else:
            sentences.append(f"Although the zone-average carbon dioxide level was {format_number(m['avg_co2'], 0)} ppm, localized high CO2 was found.{areas} This may indicate insufficient air exchange or microbiological activity in those areas. ")
    return "".join(sentences).strip()


def overall_comment(metrics: dict[str, Any]) -> str:
    zones = [metrics[z] for z in ["RMW", "PL", "FGW"]]
    ranked = sorted(zones, key=lambda m: m["risk_score"], reverse=True)
    common_counts: dict[str, int] = defaultdict(int)
    for m in zones:
        for issue in issue_names(m):
            common_counts[issue] += 1
    core_order = ["the relative humidity", "the humidity permeability rate", "the spore infiltration rate"]
    common = [issue for issue in core_order if common_counts.get(issue, 0) >= 2]
    intro = (
        f"The environmental analysis showed that the {ranked[0]['label']} had the highest environmental mold risk, "
        f"with the {ranked[1]['label']} ranking second and the {ranked[2]['label']} having the lowest risk among the three main zones. "
    )
    if common:
        intro += f"The major factors contributing to the risk were {join_list(common)}. "
    detail = ""
    if "the relative humidity" in common:
        high_zones = [m["label"] for m in zones if m["avg_rh"] and m["avg_rh"] > 60]
        detail += f"Humidity exceeded the 60% standard in {join_list(high_zones)}, indicating that multiple production or storage stages could provide enough moisture for mold germination. "
    if "the humidity permeability rate" in common:
        hp_zones = [m["label"] for m in zones if m["humidity_permeability"] and m["humidity_permeability"] > 90]
        detail += f"The humidity permeability rates in {join_list(hp_zones)} also exceeded the suggested value, meaning that outdoor humidity likely influenced the indoor environment and increased moisture persistence. "
    if "the object moisture content" in common:
        detail += "Object moisture failures showed that mold risk was not only airborne but also associated with materials, packaging, and equipment surfaces that could support local growth. "
    elif common_counts.get("the object moisture content", 0) >= 2:
        detail += "Object moisture content should still be treated as an important supporting risk because moisture failures appeared in absorbent materials, packaging, and equipment across multiple zones. "
    if "the mold spore blow index" in common:
        detail += "Low wind speed in several areas suggested that settled spores and dust may not be removed effectively from object surfaces. "
    elif common_counts.get("the mold spore blow index", 0) >= 2:
        detail += "Several individual areas also had low mold spore blow index, which could allow spores to settle on exposed surfaces. "
    conclusion = "Overall, humidity control, targeted drying of high-moisture objects, and ventilation-path improvement should be prioritized before the next audit cycle."
    return intro + detail + conclusion


def write_comments(data: dict[str, list[dict[str, Any]]], markdown_path: Path, json_path: Path) -> dict[str, Any]:
    metrics = zone_metrics(data)
    comments = {zone: zone_comment(zone, metrics[zone]) for zone in ["RMW", "PL", "FGW"]}
    comments["OVERALL"] = overall_comment(metrics)

    lines = [
        "# PYS Smart EA Environmental Comment Draft",
        "",
        "> Scope: Environment-side comments only. Microbiology comments and final risk conclusion can be added after this first automation is validated.",
        "",
        "## Metrics Used",
        "",
        "| Zone | Areas | Avg RH % | Humidity Permeability % | Avg CO2 ppm | Avg Wind m/s | Spore Infiltration % | Moisture Failed Ratio % | Risk Score |",
        "|---|---|---:|---:|---:|---:|---:|---:|---:|",
    ]
    for zone in ["RMW", "PL", "FGW"]:
        m = metrics[zone]
        lines.append(
            f"| {m['label']} | {', '.join(m['areas'])} | {format_number(m['avg_rh'])} | "
            f"{format_number(m['humidity_permeability'])} | {format_number(m['avg_co2'], 0)} | "
            f"{format_number(m['avg_wind'], 2)} | {format_number(m['spore_infiltration'])} | "
            f"{format_number(m['moisture_failed_ratio'])} | {format_number(m['risk_score'])} |"
        )
    outside = metrics["OUTSIDE"]
    lines.extend(
        [
            "",
            f"Outdoor reference: average RH {format_number(outside['avg_rh'])}%; average PM10 {format_number(outside['avg_pm10'])} μg/m³.",
            "",
            "## Draft Comments",
            "",
        ]
    )
    for zone in ["RMW", "PL", "FGW"]:
        lines.extend([f"### {metrics[zone]['label']}", "", comments[zone], ""])
    lines.extend(["### Overall Environmental Comment", "", comments["OVERALL"], ""])

    markdown_path.write_text("\n".join(lines), encoding="utf-8")
    json_path.write_text(json.dumps({"metrics": metrics, "comments": comments}, ensure_ascii=False, indent=2, default=str), encoding="utf-8")
    return {"metrics": metrics, "comments": comments}


def write_comment_basis(
    data: dict[str, list[dict[str, Any]]],
    metrics: dict[str, Any],
    output_path: Path,
    vendor_path: Path,
    target_path: Path | None,
) -> None:
    lines = [
        "# Smart EA Comment Generation Basis",
        "",
        "## Input Sources",
        "",
        f"- Vendor field-data workbook: `{vendor_path}`",
        f"- Reference / formula workbook: `{target_path}`" if target_path else "- Reference / formula workbook: not provided",
        "- Generated standardized tables: `Env. Data`, `Moisture`, `CFU`, and `Areas`.",
        "",
        "## What The Comment Script Uses",
        "",
        "- Area / warehouse grouping from the generated `Report Area` and `Main Zone` fields.",
        "- Environmental measurements: temperature, relative humidity, CO2, wind flow, and PM10.",
        "- Object moisture measurements and object classifications.",
        "- CFU colony count data when available; `CFU/m² = Count / 0.0003` is retained for manual checking.",
        "- Outdoor reference values from Outside rows for humidity permeability and spore infiltration estimates.",
        "",
        "## Comment Rules Currently Implemented",
        "",
        "- Relative humidity is flagged when the average is greater than 60%; 55-60% is treated as near-threshold.",
        "- Humidity permeability is estimated as indoor average RH / outdoor average RH * 100, and flagged when greater than 90%.",
        "- CO2 is flagged when the average is greater than 600 ppm, or when at least half of the areas are above 600 ppm.",
        "- Wind flow is flagged when the average is less than 0.4 m/s, or when at least half of the areas are below 0.4 m/s.",
        "- Spore infiltration is estimated as indoor average PM10 / outdoor average PM10 * 100, and flagged when greater than 90%.",
        "- Object moisture is flagged when measured values are greater than 10%; object groups are highlighted when their failed ratio is at least 50%.",
        "- Zone ranking uses the script's first-version environmental risk score. It is for draft ordering only and is not a replacement for the YIMS official risk result.",
        "",
        "## Current Safeguards And Limits",
        "",
        "- This script does not call an LLM. The comments are generated from deterministic Python rules and fixed English sentence templates.",
        "- The current comments cover the environment side only. Microbiology-side comments and final formal conclusions still require review.",
        "- If a value is missing, the rule does not invent a value; the output should be treated as a draft for human review.",
        "",
        "## Zone Metrics Snapshot",
        "",
        "| Zone | Avg RH % | Humidity Permeability % | Avg CO2 ppm | Avg Wind m/s | Spore Infiltration % | Moisture Failed Ratio % | Risk Score |",
        "|---|---:|---:|---:|---:|---:|---:|---:|",
    ]
    for zone in ["RMW", "PL", "FGW"]:
        m = metrics[zone]
        lines.append(
            f"| {m['label']} | {format_number(m['avg_rh'])} | "
            f"{format_number(m['humidity_permeability'])} | {format_number(m['avg_co2'], 0)} | "
            f"{format_number(m['avg_wind'], 2)} | {format_number(m['spore_infiltration'])} | "
            f"{format_number(m['moisture_failed_ratio'])} | {format_number(m['risk_score'])} |"
        )
    lines.extend(
        [
            "",
            "## Generated Table Counts",
            "",
            f"- Env. Data rows: {len(data['Env. Data'])}",
            f"- Moisture rows: {len(data['Moisture'])}",
            f"- CFU rows: {len(data['CFU'])}",
            "",
        ]
    )
    output_path.write_text("\n".join(lines), encoding="utf-8")


def write_comments_exports(comment_result: dict[str, Any], xlsx_path: Path, docx_path: Path) -> None:
    metrics = comment_result["metrics"]
    comments = comment_result["comments"]

    wb = Workbook()
    ws = wb.active
    ws.title = "Comments"
    ws.append(["Section", "Draft Comment"])
    for zone in ["RMW", "PL", "FGW"]:
        ws.append([metrics[zone]["label"], comments[zone]])
    ws.append(["Overall Environmental Comment", comments["OVERALL"]])
    style_sheet(ws)

    ws = wb.create_sheet("Metrics")
    ws.append(
        [
            "Zone",
            "Areas",
            "Avg RH %",
            "Humidity Permeability %",
            "Avg CO2 ppm",
            "Avg Wind m/s",
            "Spore Infiltration %",
            "Moisture Failed Ratio %",
            "Risk Score",
        ]
    )
    for zone in ["RMW", "PL", "FGW"]:
        m = metrics[zone]
        ws.append(
            [
                m["label"],
                ", ".join(m["areas"]),
                m["avg_rh"],
                m["humidity_permeability"],
                m["avg_co2"],
                m["avg_wind"],
                m["spore_infiltration"],
                m["moisture_failed_ratio"],
                m["risk_score"],
            ]
        )
    style_sheet(ws)
    xlsx_path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(xlsx_path)

    try:
        from docx import Document
    except ImportError:
        return

    doc = Document()
    doc.add_heading("Smart EA Environmental Comment Draft", level=1)
    doc.add_paragraph("Scope: Environment-side comments only. Microbiology comments and final risk conclusion require human review.")
    doc.add_heading("Draft Comments", level=2)
    for zone in ["RMW", "PL", "FGW"]:
        doc.add_heading(metrics[zone]["label"], level=3)
        doc.add_paragraph(comments[zone])
    doc.add_heading("Overall Environmental Comment", level=3)
    doc.add_paragraph(comments["OVERALL"])
    doc.add_heading("Metrics Used", level=2)
    table = doc.add_table(rows=1, cols=8)
    table.style = "Table Grid"
    headers = ["Zone", "Avg RH %", "HP %", "CO2 ppm", "Wind m/s", "SI %", "Moisture Failed %", "Risk Score"]
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for zone in ["RMW", "PL", "FGW"]:
        m = metrics[zone]
        cells = table.add_row().cells
        values = [
            m["label"],
            format_number(m["avg_rh"]),
            format_number(m["humidity_permeability"]),
            format_number(m["avg_co2"], 0),
            format_number(m["avg_wind"], 2),
            format_number(m["spore_infiltration"]),
            format_number(m["moisture_failed_ratio"]),
            format_number(m["risk_score"]),
        ]
        for idx, value in enumerate(values):
            cells[idx].text = value
    doc.save(docx_path)
